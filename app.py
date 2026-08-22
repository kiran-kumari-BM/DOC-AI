import os
import zipfile
import logging
import uuid
import threading

from io import BytesIO
from functools import wraps

from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_file,
    flash
)

from flask_login import (
    LoginManager,
    login_required,
    current_user
)

from werkzeug.utils import secure_filename

from docx import Document as WordDocument

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph
)

from reportlab.lib.styles import (
    getSampleStyleSheet
)

from reportlab.lib.pagesizes import A4

from config import Config

from models import (
    db,
    User,
    Document,
    ChatHistory
)

from auth import auth

from rag_engine import ask_question


# ============================================================
# APP SETUP
# ============================================================

app = Flask(__name__)

app.config.from_object(Config)


# ============================================================
# ERROR HANDLER
# ============================================================

@app.errorhandler(Exception)
def handle_exception(e):

    import traceback

    print("=" * 70, flush=True)
    print("DOC AI UNHANDLED EXCEPTION", flush=True)
    print(f"ERROR: {e}", flush=True)

    traceback.print_exc()

    print("=" * 70, flush=True)

    return (
        "Internal Server Error. Check server logs.",
        500
    )


# ============================================================
# HEALTH CHECK
# ============================================================

@app.route("/health")
def health():

    try:

        users = User.query.count()
        documents = Document.query.count()

        return {
            "status": "ok",
            "database": "connected",
            "users": users,
            "documents": documents
        }

    except Exception as e:

        import traceback

        traceback.print_exc()

        return {
            "status": "error",
            "error_type": type(e).__name__,
            "error": str(e)
        }, 500


# ============================================================
# UPLOAD FOLDER
# ============================================================

app.config["UPLOAD_FOLDER"] = os.path.join(
    os.getcwd(),
    "uploads"
)

os.makedirs(
    app.config["UPLOAD_FOLDER"],
    exist_ok=True
)


# ============================================================
# OUTPUT FOLDER
# ============================================================

app.config["OUTPUT_FOLDER"] = os.path.join(
    os.getcwd(),
    "outputs"
)

os.makedirs(
    app.config["OUTPUT_FOLDER"],
    exist_ok=True
)


# ============================================================
# DATABASE
# ============================================================

db.init_app(app)

with app.app_context():

    db.create_all()


# ============================================================
# LOGIN MANAGER
# ============================================================

login_manager = LoginManager()

login_manager.login_view = "auth.login"

login_manager.init_app(app)


# ============================================================
# AUTH BLUEPRINT
# ============================================================

app.register_blueprint(auth)


# ============================================================
# LOAD USER
# ============================================================

@login_manager.user_loader
def load_user(user_id):

    try:

        return User.query.get(
            int(user_id)
        )

    except (
        ValueError,
        TypeError
    ):

        return None


# ============================================================
# LOGGING
# ============================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s"
)


# ============================================================
# ROLE DECORATOR
# ============================================================

def role_required(role):

    def decorator(func):

        @wraps(func)
        def wrapper(*args, **kwargs):

            if (
                not current_user.is_authenticated
                or current_user.role != role
            ):

                flash(
                    "Access denied.",
                    "danger"
                )

                return redirect(
                    url_for("dashboard")
                )

            return func(
                *args,
                **kwargs
            )

        return wrapper

    return decorator


# ============================================================
# BACKGROUND OCR
# ============================================================

def process_ocr_background(
    app_instance,
    doc_id,
    path
):

    """
    Run OCR in a background thread.

    The OCR model is loaded only inside this function,
    so the Flask web process does not load the model
    during application startup.
    """

    print("=" * 70, flush=True)

    print(
        f"🔄 BACKGROUND OCR STARTED",
        flush=True
    )

    print(
        f"DOCUMENT ID: {doc_id}",
        flush=True
    )

    print(
        f"FILE: {path}",
        flush=True
    )

    print("=" * 70, flush=True)


    with app_instance.app_context():

        try:

            # ------------------------------------------------
            # GET DOCUMENT
            # ------------------------------------------------

            doc = Document.query.get(
                doc_id
            )

            if not doc:

                print(
                    f"❌ Document {doc_id} not found.",
                    flush=True
                )

                return


            # ------------------------------------------------
            # MARK PROCESSING
            # ------------------------------------------------

            doc.status = "processing"

            db.session.commit()


            # ------------------------------------------------
            # VERIFY FILE
            # ------------------------------------------------

            if not os.path.exists(path):

                raise FileNotFoundError(
                    f"Uploaded file not found: {path}"
                )


            print(
                "📄 File exists.",
                flush=True
            )


            # ------------------------------------------------
            # LOAD OCR PIPELINE
            # ------------------------------------------------

            print(
                "🧠 Importing OCR pipeline...",
                flush=True
            )

            from ocr_pipeline import run_ocr

            print(
                "✅ OCR pipeline imported.",
                flush=True
            )


            # ------------------------------------------------
            # RUN OCR
            # ------------------------------------------------

            print(
                "=" * 70,
                flush=True
            )

            print(
                f"🔍 RUNNING OCR FOR DOCUMENT {doc_id}",
                flush=True
            )

            print(
                "⚠️ First request may take several minutes "
                "because the OCR models are being loaded.",
                flush=True
            )

            print(
                "=" * 70,
                flush=True
            )

            text = run_ocr(
                path
            )


            # ------------------------------------------------
            # NORMALIZE RESULT
            # ------------------------------------------------

            if text is None:

                text = ""

            elif not isinstance(
                text,
                str
            ):

                text = str(
                    text
                )


            # ------------------------------------------------
            # SAVE RESULT
            # ------------------------------------------------

            doc.extracted_text = text

            doc.status = "completed"

            db.session.commit()


            # ------------------------------------------------
            # SUCCESS
            # ------------------------------------------------

            print("=" * 70, flush=True)

            print(
                f"✅ OCR COMPLETED",
                flush=True
            )

            print(
                f"DOCUMENT ID: {doc_id}",
                flush=True
            )

            print(
                f"EXTRACTED CHARACTERS: {len(text)}",
                flush=True
            )

            print("=" * 70, flush=True)


        except Exception as e:

            import traceback

            print("=" * 70, flush=True)

            print(
                f"❌ OCR FAILED",
                flush=True
            )

            print(
                f"DOCUMENT ID: {doc_id}",
                flush=True
            )

            print(
                f"ERROR TYPE: {type(e).__name__}",
                flush=True
            )

            print(
                f"ERROR: {str(e)}",
                flush=True
            )

            print(
                "FULL TRACEBACK:",
                flush=True
            )

            traceback.print_exc()

            print("=" * 70, flush=True)


            # ------------------------------------------------
            # UPDATE DATABASE
            # ------------------------------------------------

            try:

                doc = Document.query.get(
                    doc_id
                )

                if doc:

                    doc.status = "failed"

                    doc.extracted_text = (
                        f"OCR Error: {str(e)}"
                    )

                    db.session.commit()

            except Exception as db_error:

                print(
                    f"❌ DATABASE UPDATE FAILED: "
                    f"{db_error}",
                    flush=True
                )

                traceback.print_exc()

                db.session.rollback()


        finally:

            # ------------------------------------------------
            # CLEAN DATABASE SESSION
            # ------------------------------------------------

            try:

                db.session.remove()

            except Exception:

                pass


            print(
                f"🏁 BACKGROUND OCR FINISHED "
                f"FOR DOCUMENT {doc_id}",
                flush=True
            )


# ============================================================
# DASHBOARD
# ============================================================

@app.route("/")
@login_required
def dashboard():

    print(
        "DEBUG:",
        current_user.is_authenticated,
        current_user.id
        if current_user.is_authenticated
        else None,
        current_user.email
        if current_user.is_authenticated
        else None,
        current_user.role
        if current_user.is_authenticated
        else None,
        flush=True
    )


    # ========================================================
    # ADMIN DASHBOARD
    # ========================================================

    if current_user.role == "admin":

        users = User.query.order_by(
            User.id.desc()
        ).all()

        documents = Document.query.order_by(
            Document.id.desc()
        ).all()

        return render_template(

            "admin_dashboard.html",

            users=users,

            documents=documents,

            total_users=User.query.count(),

            total_docs=Document.query.count(),

            total_completed=
                Document.query.filter_by(
                    status="completed"
                ).count(),

            total_failed=
                Document.query.filter_by(
                    status="failed"
                ).count()

        )


    # ========================================================
    # NORMAL USER DASHBOARD
    # ========================================================

    documents = Document.query.filter_by(

        user_id=current_user.id

    ).order_by(

        Document.id.desc()

    ).all()


    return render_template(

        "dashboard.html",

        user=current_user,

        documents=documents

    )


# ============================================================
# ADMIN — DELETE USER
# ============================================================

@app.route(
    "/admin/user/<int:user_id>/delete",
    methods=["POST"]
)
@login_required
def delete_user(user_id):

    if current_user.role != "admin":

        flash(
            "❌ Access denied. Administrator privileges required.",
            "danger"
        )

        return redirect(
            url_for("dashboard")
        )


    user = User.query.get_or_404(
        user_id
    )


    if user.id == current_user.id:

        flash(
            "⚠ You cannot delete your own administrator account.",
            "warning"
        )

        return redirect(
            url_for("dashboard")
        )


    if user.role == "admin":

        flash(
            "⚠ Administrator accounts cannot be deleted.",
            "warning"
        )

        return redirect(
            url_for("dashboard")
        )


    documents = Document.query.filter_by(
        user_id=user.id
    ).all()


    for doc in documents:

        if doc.stored_path:

            try:

                if os.path.exists(
                    doc.stored_path
                ):

                    os.remove(
                        doc.stored_path
                    )

            except Exception:

                logging.exception(
                    f"Could not delete file "
                    f"for document {doc.id}"
                )

        db.session.delete(
            doc
        )


    db.session.delete(
        user
    )

    db.session.commit()


    flash(
        f"✅ User {user.email} deleted successfully.",
        "success"
    )


    return redirect(
        url_for("dashboard")
    )


# ============================================================
# ADMIN — DELETE DOCUMENT
# ============================================================

@app.route(
    "/admin/document/<int:doc_id>/delete",
    methods=["POST"]
)
@login_required
def delete_document_admin(doc_id):

    if current_user.role != "admin":

        flash(
            "❌ Access denied. Administrator privileges required.",
            "danger"
        )

        return redirect(
            url_for("dashboard")
        )


    doc = Document.query.get_or_404(
        doc_id
    )


    if doc.stored_path:

        try:

            if os.path.exists(
                doc.stored_path
            ):

                os.remove(
                    doc.stored_path
                )

        except Exception:

            logging.exception(
                f"Could not delete stored file "
                f"for document {doc.id}"
            )


    ChatHistory.query.filter_by(
        document_id=doc.id
    ).delete(
        synchronize_session=False
    )


    db.session.delete(
        doc
    )

    db.session.commit()


    flash(
        "✅ Document deleted successfully.",
        "success"
    )


    return redirect(
        url_for("dashboard")
    )


# ============================================================
# UPLOAD DOCUMENTS
# ============================================================

@app.route(
    "/upload",
    methods=["POST"]
)
@login_required
def upload():

    print("=" * 70, flush=True)

    print(
        "📥 DOC AI UPLOAD REQUEST",
        flush=True
    )

    print(
        f"USER: {current_user.email}",
        flush=True
    )

    print("=" * 70, flush=True)


    files = request.files.getlist(
        "documents"
    )


    # ========================================================
    # CHECK FILES
    # ========================================================

    if not files:

        flash(
            "No files selected.",
            "danger"
        )

        return redirect(
            url_for("dashboard")
        )


    upload_folder = app.config[
        "UPLOAD_FOLDER"
    ]


    os.makedirs(
        upload_folder,
        exist_ok=True
    )


    uploaded_count = 0


    # ========================================================
    # PROCESS EACH FILE
    # ========================================================

    for file in files:

        if not file:

            continue


        if not file.filename:

            continue


        # ----------------------------------------------------
        # SECURE FILENAME
        # ----------------------------------------------------

        original_filename = secure_filename(
            file.filename
        )


        if not original_filename:

            continue


        # ----------------------------------------------------
        # UNIQUE FILENAME
        # ----------------------------------------------------

        unique_name = (
            f"{uuid.uuid4()}_"
            f"{original_filename}"
        )


        save_path = os.path.join(
            upload_folder,
            unique_name
        )


        # ====================================================
        # SAVE FILE
        # ====================================================

        try:

            file.save(
                save_path
            )

            print(
                f"✅ FILE SAVED: {save_path}",
                flush=True
            )

        except Exception:

            logging.exception(
                "Failed to save uploaded file."
            )

            flash(
                f"Could not save {original_filename}.",
                "danger"
            )

            continue


        # ====================================================
        # CREATE DATABASE RECORD
        # ====================================================

        try:

            doc = Document(

                filename=unique_name,

                stored_path=save_path,

                user_id=current_user.id,

                status="processing"

            )


            db.session.add(
                doc
            )


            db.session.commit()


            uploaded_count += 1


            print(
                f"✅ DATABASE RECORD CREATED",
                flush=True
            )

            print(
                f"DOCUMENT ID: {doc.id}",
                flush=True
            )

        except Exception:

            db.session.rollback()

            logging.exception(
                "Failed to create document record."
            )


            try:

                if os.path.exists(
                    save_path
                ):

                    os.remove(
                        save_path
                    )

            except Exception:

                pass


            continue


        # ====================================================
        # START BACKGROUND OCR
        # ====================================================

        try:

            print(
                f"🚀 STARTING BACKGROUND OCR "
                f"FOR DOCUMENT {doc.id}",
                flush=True
            )


            thread = threading.Thread(

                target=process_ocr_background,

                args=(
                    app,
                    doc.id,
                    save_path
                ),

                daemon=True

            )


            thread.start()


            print(
                f"✅ BACKGROUND OCR THREAD STARTED "
                f"FOR DOCUMENT {doc.id}",
                flush=True
            )


        except Exception as e:

            logging.exception(
                f"❌ Failed to start OCR "
                f"for document {doc.id}"
            )


            doc.status = "failed"


            doc.extracted_text = (
                "OCR task could not be started: "
                f"{str(e)}"
            )


            db.session.commit()


    # ========================================================
    # RESULT
    # ========================================================

    if uploaded_count > 0:

        flash(

            f"{uploaded_count} document(s) uploaded. "
            "OCR processing started.",

            "success"

        )

    else:

        flash(
            "No documents were uploaded.",
            "warning"
        )


    return redirect(
        url_for("dashboard")
    )


# ============================================================
# VIEW DOCUMENT
# ============================================================

@app.route(
    "/document/<int:doc_id>",
    methods=["GET", "POST"]
)
@login_required
def view_document(doc_id):

    doc = Document.query.get_or_404(
        doc_id
    )


    # --------------------------------------------------------
    # SECURITY
    # --------------------------------------------------------

    if (
        doc.user_id != current_user.id
        and current_user.role != "admin"
    ):

        return "Unauthorized", 403


    # --------------------------------------------------------
    # CHAT HISTORY
    # --------------------------------------------------------

    chat_history = ChatHistory.query.filter_by(

        document_id=doc.id

    ).order_by(

        ChatHistory.id.asc()

    ).all()


    # --------------------------------------------------------
    # SAVE EDITED OCR TEXT
    # --------------------------------------------------------

    if request.method == "POST":

        edited_text = request.form.get(
            "edited_text",
            ""
        )


        doc.extracted_text = edited_text


        db.session.commit()


        flash(
            "Document updated successfully.",
            "success"
        )


        return redirect(

            url_for(
                "view_document",
                doc_id=doc.id
            )

        )


    # --------------------------------------------------------
    # RENDER
    # --------------------------------------------------------

    return render_template(

        "view_document.html",

        doc=doc,

        chat_history=chat_history

    )


# ============================================================
# SERVE ORIGINAL DOCUMENT
# ============================================================

@app.route(
    "/document/file/<int:doc_id>"
)
@login_required
def document_file(doc_id):

    doc = Document.query.get_or_404(
        doc_id
    )


    # --------------------------------------------------------
    # SECURITY
    # --------------------------------------------------------

    if (
        doc.user_id != current_user.id
        and current_user.role != "admin"
    ):

        return "Unauthorized", 403


    if not doc.stored_path:

        return "File not found", 404


    if not os.path.exists(
        doc.stored_path
    ):

        logging.error(
            f"Stored file missing: "
            f"{doc.stored_path}"
        )

        return "File not found", 404


    return send_file(
        doc.stored_path
    )


# ============================================================
# DELETE DOCUMENT
# ============================================================

@app.route(
    "/document/delete/<int:doc_id>",
    methods=["GET"]
)
@login_required
def delete_document_user(doc_id):

    doc = Document.query.get_or_404(
        doc_id
    )


    # --------------------------------------------------------
    # SECURITY
    # --------------------------------------------------------

    if (
        doc.user_id != current_user.id
        and current_user.role != "admin"
    ):

        return "Unauthorized", 403


    # --------------------------------------------------------
    # DELETE FILE
    # --------------------------------------------------------

    if doc.stored_path:

        try:

            if os.path.exists(
                doc.stored_path
            ):

                os.remove(
                    doc.stored_path
                )

                logging.info(
                    f"Deleted file: "
                    f"{doc.stored_path}"
                )

        except Exception as e:

            logging.error(
                f"Could not delete file: {e}"
            )


    # --------------------------------------------------------
    # DELETE CHAT HISTORY
    # --------------------------------------------------------

    ChatHistory.query.filter_by(

        document_id=doc.id

    ).delete(

        synchronize_session=False

    )


    # --------------------------------------------------------
    # DELETE DATABASE RECORD
    # --------------------------------------------------------

    db.session.delete(
        doc
    )

    db.session.commit()


    logging.info(
        f"Document {doc_id} deleted by "
        f"user {current_user.id}"
    )


    flash(
        "Document deleted successfully.",
        "success"
    )


    return redirect(
        url_for("dashboard")
    )


# ============================================================
# ASK DOC AI
# ============================================================

@app.route(
    "/ask/<int:doc_id>",
    methods=["POST"]
)
@login_required
def ask_route(doc_id):

    doc = Document.query.get_or_404(
        doc_id
    )


    # --------------------------------------------------------
    # SECURITY
    # --------------------------------------------------------

    if (
        doc.user_id != current_user.id
        and current_user.role != "admin"
    ):

        return "Unauthorized", 403


    question = request.form.get(
        "question",
        ""
    ).strip()


    if not question:

        flash(
            "Please enter a question.",
            "warning"
        )

        return redirect(

            url_for(
                "view_document",
                doc_id=doc.id
            )

        )


    if not doc.extracted_text:

        flash(
            "No extracted text is available.",
            "warning"
        )

        return redirect(

            url_for(
                "view_document",
                doc_id=doc.id
            )

        )


    # --------------------------------------------------------
    # RAG
    # --------------------------------------------------------

    try:

        answer, citation = ask_question(

            doc.extracted_text,

            question

        )


        chat = ChatHistory(

            document_id=doc.id,

            question=question,

            answer=answer,

            citation=citation

        )


        db.session.add(
            chat
        )


        db.session.commit()


    except Exception as e:

        db.session.rollback()


        logging.exception(
            "RAG processing failed"
        )


        print(
            "RAG ERROR:",
            e,
            flush=True
        )


        flash(
            "Error processing your question.",
            "danger"
        )


    return redirect(

        url_for(
            "view_document",
            doc_id=doc.id
        )

    )


# ============================================================
# DOWNLOAD PDF
# ============================================================

@app.route(
    "/download/pdf/<int:doc_id>"
)
@login_required
def download_pdf(doc_id):

    doc = Document.query.get_or_404(
        doc_id
    )


    # --------------------------------------------------------
    # SECURITY
    # --------------------------------------------------------

    if (
        doc.user_id != current_user.id
        and current_user.role != "admin"
    ):

        return "Unauthorized", 403


    output_path = os.path.join(

        app.config["OUTPUT_FOLDER"],

        f"{doc.id}.pdf"

    )


    pdf = SimpleDocTemplate(

        output_path,

        pagesize=A4

    )


    styles = getSampleStyleSheet()


    elements = []


    for line in (

        doc.extracted_text or ""

    ).split("\n"):

        safe_line = (

            line
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")

        )


        elements.append(

            Paragraph(

                safe_line,

                styles["Normal"]

            )

        )


    pdf.build(
        elements
    )


    return send_file(

        output_path,

        as_attachment=True,

        download_name=
            f"DOC_AI_{doc.id}.pdf"

    )


# ============================================================
# DOWNLOAD WORD
# ============================================================

@app.route(
    "/download/word/<int:doc_id>"
)
@login_required
def download_word(doc_id):

    doc = Document.query.get_or_404(
        doc_id
    )


    # --------------------------------------------------------
    # SECURITY
    # --------------------------------------------------------

    if (
        doc.user_id != current_user.id
        and current_user.role != "admin"
    ):

        return "Unauthorized", 403


    output_path = os.path.join(

        app.config["OUTPUT_FOLDER"],

        f"{doc.id}.docx"

    )


    word_doc = WordDocument()


    word_doc.add_paragraph(

        doc.extracted_text or ""

    )


    word_doc.save(
        output_path
    )


    return send_file(

        output_path,

        as_attachment=True,

        download_name=
            f"DOC_AI_{doc.id}.docx"

    )


# ============================================================
# BULK DOWNLOAD
# ============================================================

@app.route(
    "/download/bulk",
    methods=["POST"]
)
@login_required
def bulk_download():

    selected_ids = request.form.getlist(
        "selected_docs"
    )


    if not selected_ids:

        flash(
            "No documents selected.",
            "warning"
        )

        return redirect(
            url_for("dashboard")
        )


    memory_file = BytesIO()


    with zipfile.ZipFile(

        memory_file,

        "w",

        zipfile.ZIP_DEFLATED

    ) as zf:


        for doc_id in selected_ids:

            try:

                doc = Document.query.get(
                    int(doc_id)
                )

            except (
                ValueError,
                TypeError
            ):

                continue


            if not doc:

                continue


            # ------------------------------------------------
            # SECURITY
            # ------------------------------------------------

            if (
                doc.user_id != current_user.id
                and current_user.role != "admin"
            ):

                continue


            # ------------------------------------------------
            # SAFE NAME
            # ------------------------------------------------

            safe_name = secure_filename(
                doc.filename
            )


            if not safe_name:

                safe_name = (
                    f"document_{doc.id}"
                )


            # ------------------------------------------------
            # ADD TO ZIP
            # ------------------------------------------------

            zf.writestr(

                f"{safe_name}.txt",

                doc.extracted_text or ""

            )


    memory_file.seek(0)


    return send_file(

        memory_file,

        download_name=
            "DOC_AI_documents.zip",

        as_attachment=True

    )


# ============================================================
# RUN
# ============================================================

if __name__ == "__main__":

    with app.app_context():

        db.create_all()


    app.run(

        host="0.0.0.0",

        port=int(

            os.environ.get(
                "PORT",
                5000
            )

        ),

        debug=False

    )